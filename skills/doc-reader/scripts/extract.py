#!/usr/bin/env python3
"""
Document Intelligence Extractor

Universal document text/table/metadata extraction supporting:
- PDF (native text, OCR fallback, scanned documents)
- DOCX (Word documents with tables, images, headers/footers)
- XLSX/CSV (spreadsheets with multi-sheet support)
- Images (OCR via Tesseract)
- HTML (via BeautifulSoup)
- Plain text / Markdown

Extraction modes:
- text:    Full text extraction (default)
- tables:  Table extraction as CSV/JSON/markdown
- meta:    Document metadata (author, dates, pages, etc.)
- summary: Text + tables + metadata combined
- ocr:     Force OCR even on native-text PDFs
- pages:   Extract specific page ranges

Usage:
  python3 extract.py <file> [--mode text|tables|meta|summary|ocr|pages]
                             [--pages 1-5,8,10-12]
                             [--table-format csv|json|markdown]
                             [--output <file>]
                             [--ocr-lang eng+nor]
                             [--max-pages 50]
                             [--verbose]
"""

import argparse
import csv
import io
import json
import logging
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Optional

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ==================== PDF EXTRACTION ====================

class PDFExtractor:
    """Multi-strategy PDF text and table extraction."""

    def __init__(self, filepath: str, ocr_lang: str = "eng"):
        self.filepath = filepath
        self.ocr_lang = ocr_lang
        self._doc = None

    def _get_doc(self):
        if self._doc is None:
            import pymupdf
            self._doc = pymupdf.open(self.filepath)
        return self._doc

    @property
    def page_count(self) -> int:
        return self._get_doc().page_count

    def metadata(self) -> dict:
        doc = self._get_doc()
        meta = doc.metadata or {}
        return {
            "title": meta.get("title", ""),
            "author": meta.get("author", ""),
            "subject": meta.get("subject", ""),
            "creator": meta.get("creator", ""),
            "producer": meta.get("producer", ""),
            "creation_date": meta.get("creationDate", ""),
            "mod_date": meta.get("modDate", ""),
            "pages": doc.page_count,
            "file_size": os.path.getsize(self.filepath),
            "file_size_human": _human_size(os.path.getsize(self.filepath)),
            "encrypted": doc.is_encrypted,
        }

    def extract_text(self, pages: list = None) -> str:
        """Extract text using PyMuPDF (best for native text PDFs)."""
        doc = self._get_doc()
        page_range = pages or range(doc.page_count)
        texts = []

        for i in page_range:
            if i >= doc.page_count:
                continue
            page = doc[i]
            text = page.get_text("text")

            # Check if page has meaningful text
            if text and len(text.strip()) > 20:
                texts.append(f"--- Page {i + 1} ---\n{text.strip()}")
            else:
                # Might be scanned — try OCR
                ocr_text = self._ocr_page(page, i)
                if ocr_text:
                    texts.append(f"--- Page {i + 1} (OCR) ---\n{ocr_text.strip()}")
                else:
                    texts.append(f"--- Page {i + 1} ---\n[No extractable text]")

        return "\n\n".join(texts)

    def extract_text_pdftotext(self, pages: list = None) -> str:
        """Extract text using pdftotext (poppler) — better layout preservation."""
        cmd = ["pdftotext", "-layout"]
        if pages:
            # pdftotext uses -f (first) and -l (last)
            cmd.extend(["-f", str(min(pages) + 1), "-l", str(max(pages) + 1)])
        cmd.extend([self.filepath, "-"])

        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            return result.stdout
        except (subprocess.TimeoutExpired, FileNotFoundError):
            return ""

    def extract_text_ocr(self, pages: list = None) -> str:
        """Force OCR on all pages (for scanned documents)."""
        doc = self._get_doc()
        page_range = pages or range(doc.page_count)
        texts = []

        for i in page_range:
            if i >= doc.page_count:
                continue
            page = doc[i]
            ocr_text = self._ocr_page(page, i)
            texts.append(f"--- Page {i + 1} (OCR) ---\n{(ocr_text or '[OCR failed]').strip()}")

        return "\n\n".join(texts)

    def _ocr_page(self, page, page_num: int) -> str:
        """OCR a single page by rendering to image and running Tesseract."""
        try:
            # Render page at 300 DPI
            pix = page.get_pixmap(dpi=300)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                pix.save(f.name)
                tmp_path = f.name

            try:
                result = subprocess.run(
                    ["tesseract", tmp_path, "stdout", "-l", self.ocr_lang, "--psm", "6"],
                    capture_output=True, text=True, timeout=30
                )
                return result.stdout
            finally:
                os.unlink(tmp_path)
        except Exception as e:
            logger.debug(f"OCR failed for page {page_num}: {e}")
            return ""

    def extract_tables(self, pages: list = None, fmt: str = "markdown") -> list:
        """Extract tables using pdfplumber (best general-purpose table extractor)."""
        import pdfplumber

        tables_out = []
        with pdfplumber.open(self.filepath) as pdf:
            page_range = pages or range(len(pdf.pages))

            for i in page_range:
                if i >= len(pdf.pages):
                    continue
                page = pdf.pages[i]
                tables = page.extract_tables()

                for t_idx, table in enumerate(tables):
                    if not table or not any(any(cell for cell in row) for row in table):
                        continue

                    # Clean cells
                    cleaned = []
                    for row in table:
                        cleaned.append([
                            (cell or "").strip().replace("\n", " ")
                            for cell in row
                        ])

                    table_data = {
                        "page": i + 1,
                        "table_index": t_idx + 1,
                        "rows": len(cleaned),
                        "cols": max(len(r) for r in cleaned) if cleaned else 0,
                        "data": cleaned,
                    }

                    if fmt == "markdown":
                        table_data["formatted"] = _table_to_markdown(cleaned)
                    elif fmt == "csv":
                        table_data["formatted"] = _table_to_csv(cleaned)
                    elif fmt == "json":
                        # Use first row as headers if it looks like a header
                        if len(cleaned) > 1:
                            headers = cleaned[0]
                            rows = [dict(zip(headers, row)) for row in cleaned[1:]]
                            table_data["formatted"] = json.dumps(rows, indent=2)
                        else:
                            table_data["formatted"] = json.dumps(cleaned, indent=2)

                    tables_out.append(table_data)

        return tables_out

    def extract_tables_tabula(self, pages: list = None) -> list:
        """Extract tables using tabula-java (better for bordered tables)."""
        try:
            import tabula

            page_str = "all"
            if pages:
                page_str = ",".join(str(p + 1) for p in pages)

            dfs = tabula.read_pdf(
                self.filepath,
                pages=page_str,
                multiple_tables=True,
                silent=True,
            )

            tables_out = []
            for i, df in enumerate(dfs):
                if df.empty:
                    continue
                tables_out.append({
                    "table_index": i + 1,
                    "rows": len(df),
                    "cols": len(df.columns),
                    "columns": list(df.columns),
                    "data_csv": df.to_csv(index=False),
                    "data_markdown": df.to_markdown(index=False),
                })
            return tables_out
        except Exception as e:
            logger.warning(f"tabula extraction failed: {e}")
            return []

    def extract_images(self, pages: list = None, output_dir: str = None) -> list:
        """Extract embedded images from PDF."""
        doc = self._get_doc()
        page_range = pages or range(doc.page_count)
        images = []

        out_dir = Path(output_dir) if output_dir else Path(tempfile.mkdtemp(prefix="doc_images_"))
        out_dir.mkdir(parents=True, exist_ok=True)

        for i in page_range:
            if i >= doc.page_count:
                continue
            page = doc[i]
            for img_idx, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                try:
                    pix = pymupdf.Pixmap(doc, xref)
                    if pix.n > 4:  # CMYK
                        pix = pymupdf.Pixmap(pymupdf.csRGB, pix)

                    ext = "png" if pix.alpha else "png"
                    img_path = out_dir / f"page{i + 1}_img{img_idx + 1}.{ext}"
                    pix.save(str(img_path))
                    images.append({
                        "page": i + 1,
                        "index": img_idx + 1,
                        "path": str(img_path),
                        "width": pix.width,
                        "height": pix.height,
                    })
                except Exception as e:
                    logger.debug(f"Failed to extract image {xref} from page {i + 1}: {e}")

        return images

    def close(self):
        if self._doc:
            self._doc.close()
            self._doc = None


# ==================== DOCX EXTRACTION ====================

class DOCXExtractor:
    """Microsoft Word document extraction."""

    def __init__(self, filepath: str):
        self.filepath = filepath

    def metadata(self) -> dict:
        from docx import Document
        doc = Document(self.filepath)
        props = doc.core_properties
        return {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created or ""),
            "modified": str(props.modified or ""),
            "last_modified_by": props.last_modified_by or "",
            "revision": props.revision or 0,
            "category": props.category or "",
            "comments": props.comments or "",
            "file_size": os.path.getsize(self.filepath),
            "file_size_human": _human_size(os.path.getsize(self.filepath)),
        }

    def extract_text(self) -> str:
        from docx import Document
        doc = Document(self.filepath)

        texts = []

        for element in doc.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':
                # Paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para:
                    style = para.style.name if para.style else ""
                    text = para.text.strip()
                    if not text:
                        continue

                    # Format headings
                    if "Heading 1" in style:
                        texts.append(f"\n# {text}")
                    elif "Heading 2" in style:
                        texts.append(f"\n## {text}")
                    elif "Heading 3" in style:
                        texts.append(f"\n### {text}")
                    elif "List" in style:
                        texts.append(f"  • {text}")
                    else:
                        texts.append(text)

            elif tag == 'tbl':
                # Table
                for table in doc.tables:
                    if table._element is element:
                        rows = []
                        for row in table.rows:
                            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                            rows.append(cells)
                        if rows:
                            texts.append("\n" + _table_to_markdown(rows) + "\n")
                        break

        return "\n".join(texts)

    def extract_tables(self, fmt: str = "markdown") -> list:
        from docx import Document
        doc = Document(self.filepath)
        tables_out = []

        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)

            if not rows or not any(any(c for c in row) for row in rows):
                continue

            table_data = {
                "table_index": t_idx + 1,
                "rows": len(rows),
                "cols": max(len(r) for r in rows) if rows else 0,
                "data": rows,
            }

            if fmt == "markdown":
                table_data["formatted"] = _table_to_markdown(rows)
            elif fmt == "csv":
                table_data["formatted"] = _table_to_csv(rows)
            elif fmt == "json":
                if len(rows) > 1:
                    headers = rows[0]
                    records = [dict(zip(headers, row)) for row in rows[1:]]
                    table_data["formatted"] = json.dumps(records, indent=2)
                else:
                    table_data["formatted"] = json.dumps(rows, indent=2)

            tables_out.append(table_data)

        return tables_out


# ==================== SPREADSHEET EXTRACTION ====================

class SpreadsheetExtractor:
    """Excel/CSV extraction with multi-sheet support."""

    def __init__(self, filepath: str):
        self.filepath = filepath
        self.ext = Path(filepath).suffix.lower()

    def metadata(self) -> dict:
        if self.ext == '.csv':
            return {
                "format": "CSV",
                "file_size": os.path.getsize(self.filepath),
                "file_size_human": _human_size(os.path.getsize(self.filepath)),
            }

        from openpyxl import load_workbook
        wb = load_workbook(self.filepath, read_only=True, data_only=True)
        return {
            "format": "XLSX",
            "sheets": wb.sheetnames,
            "sheet_count": len(wb.sheetnames),
            "file_size": os.path.getsize(self.filepath),
            "file_size_human": _human_size(os.path.getsize(self.filepath)),
        }

    def extract_text(self, sheet: str = None) -> str:
        if self.ext == '.csv':
            return self._extract_csv()
        return self._extract_xlsx(sheet)

    def _extract_csv(self) -> str:
        with open(self.filepath, 'r', newline='', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f)
            rows = list(reader)
        if rows:
            return _table_to_markdown(rows)
        return ""

    def _extract_xlsx(self, sheet: str = None) -> str:
        from openpyxl import load_workbook
        wb = load_workbook(self.filepath, read_only=True, data_only=True)

        sheets_to_process = [sheet] if sheet else wb.sheetnames
        texts = []

        for sheet_name in sheets_to_process:
            if sheet_name not in wb.sheetnames:
                continue
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c for c in cells):
                    rows.append(cells)

            if rows:
                texts.append(f"## Sheet: {sheet_name}\n")
                texts.append(_table_to_markdown(rows))
                texts.append("")

        return "\n".join(texts)

    def extract_tables(self, sheet: str = None, fmt: str = "markdown") -> list:
        if self.ext == '.csv':
            with open(self.filepath, 'r', newline='', encoding='utf-8', errors='replace') as f:
                rows = list(csv.reader(f))
            if rows:
                return [{
                    "sheet": "csv",
                    "table_index": 1,
                    "rows": len(rows),
                    "cols": max(len(r) for r in rows),
                    "data": rows,
                    "formatted": _table_to_markdown(rows) if fmt == "markdown" else _table_to_csv(rows),
                }]
            return []

        from openpyxl import load_workbook
        wb = load_workbook(self.filepath, read_only=True, data_only=True)
        tables = []
        sheets_to_process = [sheet] if sheet else wb.sheetnames

        for sheet_name in sheets_to_process:
            if sheet_name not in wb.sheetnames:
                continue
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c for c in cells):
                    rows.append(cells)
            if rows:
                table_data = {
                    "sheet": sheet_name,
                    "table_index": 1,
                    "rows": len(rows),
                    "cols": max(len(r) for r in rows),
                    "data": rows,
                }
                if fmt == "markdown":
                    table_data["formatted"] = _table_to_markdown(rows)
                elif fmt == "csv":
                    table_data["formatted"] = _table_to_csv(rows)
                elif fmt == "json":
                    headers = rows[0]
                    records = [dict(zip(headers, r)) for r in rows[1:]]
                    table_data["formatted"] = json.dumps(records, indent=2)
                tables.append(table_data)

        return tables


# ==================== IMAGE OCR EXTRACTION ====================

class ImageExtractor:
    """OCR extraction from images (PNG, JPG, TIFF, BMP)."""

    def __init__(self, filepath: str, lang: str = "eng"):
        self.filepath = filepath
        self.lang = lang

    def metadata(self) -> dict:
        from PIL import Image
        img = Image.open(self.filepath)
        return {
            "format": img.format,
            "mode": img.mode,
            "width": img.width,
            "height": img.height,
            "file_size": os.path.getsize(self.filepath),
            "file_size_human": _human_size(os.path.getsize(self.filepath)),
        }

    def extract_text(self) -> str:
        try:
            result = subprocess.run(
                ["tesseract", self.filepath, "stdout", "-l", self.lang, "--psm", "6"],
                capture_output=True, text=True, timeout=60
            )
            return result.stdout.strip()
        except Exception as e:
            return f"[OCR failed: {e}]"


# ==================== HTML EXTRACTION ====================

class HTMLExtractor:
    """HTML to clean text extraction."""

    def __init__(self, filepath: str):
        self.filepath = filepath

    def extract_text(self) -> str:
        from bs4 import BeautifulSoup
        with open(self.filepath, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f, 'html.parser')

        # Remove script and style
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()

        return soup.get_text(separator='\n', strip=True)

    def metadata(self) -> dict:
        from bs4 import BeautifulSoup
        with open(self.filepath, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f, 'html.parser')

        title = soup.find('title')
        return {
            "title": title.get_text(strip=True) if title else "",
            "file_size": os.path.getsize(self.filepath),
        }


# ==================== HELPERS ====================

def _table_to_markdown(rows: list) -> str:
    """Convert 2D list to markdown table."""
    if not rows:
        return ""

    # Normalize column count
    max_cols = max(len(r) for r in rows)
    normalized = [r + [""] * (max_cols - len(r)) for r in rows]

    # Calculate column widths
    widths = [
        max(len(str(normalized[r][c])) for r in range(len(normalized)))
        for c in range(max_cols)
    ]
    widths = [max(w, 3) for w in widths]

    lines = []
    for i, row in enumerate(normalized):
        cells = [str(cell).ljust(widths[j]) for j, cell in enumerate(row)]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("| " + " | ".join("-" * w for w in widths) + " |")

    return "\n".join(lines)


def _table_to_csv(rows: list) -> str:
    """Convert 2D list to CSV string."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerows(rows)
    return output.getvalue()


def _human_size(size_bytes: int) -> str:
    """Convert bytes to human-readable size."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def _parse_page_range(page_str: str, max_pages: int) -> list:
    """Parse '1-5,8,10-12' into [0,1,2,3,4,7,9,10,11]."""
    pages = set()
    for part in page_str.split(","):
        part = part.strip()
        if "-" in part:
            start, end = part.split("-", 1)
            start = max(0, int(start) - 1)
            end = min(max_pages, int(end))
            pages.update(range(start, end))
        else:
            p = int(part) - 1
            if 0 <= p < max_pages:
                pages.add(p)
    return sorted(pages)


def _detect_format(filepath: str) -> str:
    """Detect document format from extension."""
    ext = Path(filepath).suffix.lower()
    format_map = {
        '.pdf': 'pdf',
        '.docx': 'docx', '.doc': 'docx',
        '.xlsx': 'xlsx', '.xls': 'xlsx',
        '.csv': 'csv',
        '.png': 'image', '.jpg': 'image', '.jpeg': 'image',
        '.tiff': 'image', '.tif': 'image', '.bmp': 'image',
        '.html': 'html', '.htm': 'html',
        '.txt': 'text', '.md': 'text', '.log': 'text',
        '.json': 'text', '.xml': 'text', '.yaml': 'text', '.yml': 'text',
    }
    return format_map.get(ext, 'unknown')


# ==================== MAIN EXTRACTION LOGIC ====================

def extract_document(filepath: str, mode: str = "text", pages: str = None,
                     table_format: str = "markdown", ocr_lang: str = "eng",
                     max_pages: int = 200, sheet: str = None,
                     verbose: bool = False) -> dict:
    """
    Main entry point. Returns a dict with:
    - format: detected format
    - mode: extraction mode used
    - metadata: document metadata
    - text: extracted text (if applicable)
    - tables: extracted tables (if applicable)
    - images: extracted images (if applicable)
    - stats: extraction statistics
    """
    filepath = str(Path(filepath).expanduser().resolve())

    if not os.path.exists(filepath):
        return {"error": f"File not found: {filepath}"}

    doc_format = _detect_format(filepath)
    result = {
        "file": filepath,
        "format": doc_format,
        "mode": mode,
    }

    if verbose:
        logger.setLevel(logging.DEBUG)

    # ---- PDF ----
    if doc_format == 'pdf':
        extractor = PDFExtractor(filepath, ocr_lang=ocr_lang)
        result["metadata"] = extractor.metadata()

        page_list = None
        if pages:
            page_list = _parse_page_range(pages, extractor.page_count)
        elif extractor.page_count > max_pages:
            logger.warning(f"Document has {extractor.page_count} pages, limiting to {max_pages}")
            page_list = list(range(max_pages))

        if mode in ("text", "summary"):
            result["text"] = extractor.extract_text(page_list)
        elif mode == "ocr":
            result["text"] = extractor.extract_text_ocr(page_list)
        elif mode == "layout":
            result["text"] = extractor.extract_text_pdftotext(page_list)

        if mode in ("tables", "summary"):
            result["tables"] = extractor.extract_tables(page_list, fmt=table_format)

        if mode == "images":
            result["images"] = extractor.extract_images(page_list)

        if mode == "meta":
            pass  # metadata already set

        extractor.close()

    # ---- DOCX ----
    elif doc_format == 'docx':
        extractor = DOCXExtractor(filepath)
        result["metadata"] = extractor.metadata()

        if mode in ("text", "summary"):
            result["text"] = extractor.extract_text()

        if mode in ("tables", "summary"):
            result["tables"] = extractor.extract_tables(fmt=table_format)

    # ---- Spreadsheet ----
    elif doc_format in ('xlsx', 'csv'):
        extractor = SpreadsheetExtractor(filepath)
        result["metadata"] = extractor.metadata()

        if mode in ("text", "summary"):
            result["text"] = extractor.extract_text(sheet=sheet)

        if mode in ("tables", "summary"):
            result["tables"] = extractor.extract_tables(sheet=sheet, fmt=table_format)

    # ---- Image ----
    elif doc_format == 'image':
        extractor = ImageExtractor(filepath, lang=ocr_lang)
        result["metadata"] = extractor.metadata()
        result["text"] = extractor.extract_text()

    # ---- HTML ----
    elif doc_format == 'html':
        extractor = HTMLExtractor(filepath)
        result["metadata"] = extractor.metadata()
        result["text"] = extractor.extract_text()

    # ---- Text ----
    elif doc_format == 'text':
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            result["text"] = f.read()
        result["metadata"] = {
            "file_size": os.path.getsize(filepath),
            "file_size_human": _human_size(os.path.getsize(filepath)),
            "lines": result["text"].count("\n") + 1,
        }

    else:
        result["error"] = f"Unsupported format: {doc_format} ({Path(filepath).suffix})"

    # Stats
    stats = {}
    if "text" in result and result["text"]:
        text = result["text"]
        stats["chars"] = len(text)
        stats["words"] = len(text.split())
        stats["lines"] = text.count("\n") + 1
    if "tables" in result:
        stats["tables_found"] = len(result["tables"])
        stats["total_table_rows"] = sum(t.get("rows", 0) for t in result["tables"])
    if "images" in result:
        stats["images_found"] = len(result["images"])
    result["stats"] = stats

    return result


# ==================== CLI ====================

def main():
    parser = argparse.ArgumentParser(
        description="Document Intelligence Extractor — PDF, DOCX, XLSX, Images, HTML",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s report.pdf                           # Extract text from PDF
  %(prog)s report.pdf --mode tables             # Extract tables only
  %(prog)s report.pdf --mode summary            # Text + tables + metadata
  %(prog)s report.pdf --mode ocr                # Force OCR (scanned docs)
  %(prog)s report.pdf --pages 1-5,8             # Specific pages
  %(prog)s contract.docx                        # Extract Word doc
  %(prog)s data.xlsx --sheet "Sheet1"           # Specific Excel sheet
  %(prog)s scan.png --ocr-lang eng+nor          # OCR image (English + Norwegian)
  %(prog)s report.pdf --mode tables --table-format json --output tables.json
        """
    )
    parser.add_argument('file', help='Document file path')
    parser.add_argument('--mode', choices=['text', 'tables', 'meta', 'summary', 'ocr', 'layout', 'images'],
                        default='text', help='Extraction mode (default: text)')
    parser.add_argument('--pages', type=str, help='Page range: 1-5,8,10-12')
    parser.add_argument('--table-format', choices=['markdown', 'csv', 'json'], default='markdown',
                        help='Table output format (default: markdown)')
    parser.add_argument('--output', '-o', type=str, help='Write output to file')
    parser.add_argument('--ocr-lang', type=str, default='eng', help='Tesseract OCR language(s) (default: eng)')
    parser.add_argument('--max-pages', type=int, default=200, help='Max pages to process (default: 200)')
    parser.add_argument('--sheet', type=str, help='Excel sheet name')
    parser.add_argument('--json', action='store_true', help='Output as JSON')
    parser.add_argument('--verbose', '-v', action='store_true', help='Verbose logging')
    args = parser.parse_args()

    result = extract_document(
        filepath=args.file,
        mode=args.mode,
        pages=args.pages,
        table_format=args.table_format,
        ocr_lang=args.ocr_lang,
        max_pages=args.max_pages,
        sheet=args.sheet,
        verbose=args.verbose,
    )

    if args.json:
        output = json.dumps(result, indent=2, default=str)
    else:
        # Human-readable output
        parts = []

        if result.get("error"):
            parts.append(f"ERROR: {result['error']}")

        if result.get("metadata") and args.mode in ("meta", "summary"):
            parts.append("=== METADATA ===")
            for k, v in result["metadata"].items():
                if v:
                    parts.append(f"  {k}: {v}")

        if result.get("text"):
            if args.mode in ("summary",):
                parts.append("\n=== TEXT ===")
            parts.append(result["text"])

        if result.get("tables"):
            parts.append("\n=== TABLES ===")
            for t in result["tables"]:
                label_parts = [f"Table {t.get('table_index', '?')}"]
                if t.get('page'):
                    label_parts.append(f"Page {t['page']}")
                if t.get('sheet'):
                    label_parts.append(f"Sheet: {t['sheet']}")
                label_parts.append(f"{t['rows']}×{t['cols']}")
                parts.append(f"\n--- {' | '.join(label_parts)} ---")
                parts.append(t.get("formatted", ""))

        if result.get("images"):
            parts.append("\n=== IMAGES ===")
            for img in result["images"]:
                parts.append(f"  Page {img['page']}, Image {img['index']}: "
                           f"{img['width']}×{img['height']} → {img['path']}")

        if result.get("stats"):
            parts.append(f"\n=== STATS ===")
            for k, v in result["stats"].items():
                parts.append(f"  {k}: {v}")

        output = "\n".join(parts)

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(output)
        print(f"Written to {args.output}")
    else:
        print(output)


if __name__ == '__main__':
    main()
